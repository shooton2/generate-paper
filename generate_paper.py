import app
from openai import OpenAI
from docx import Document
client = OpenAI(api_key="sk-uibaSqnGRtJXAIybs1q5zUp2TrnBdiLo7gUdE1tbuKQYFrce",base_url="https://api.chatanywhere.tech/v1")


###用户输入关键词，通过与用户之间交互，生成一篇论文
class Word2paper:
    def __init__(self,user_input,user_satistied_problem=False,user_satisfied_keywords=False,user_satisfied_method=False,num_papers=5):  ##用户输入指定信息
        self.user_satisfied_problem = user_satistied_problem ##用户是否满意给出的问题
        self.user_satisfied_keywords = user_satisfied_keywords ##用户是否满意给出的关键词用于搜索semantic scholar
        self.user_satisfied_method = user_satisfied_method ##用户是否满意给出的creative method
        self.prompts = app.load_all_prompts()
        self.history = []
        self.prompt_enhance = self.prompts["prompt_enhance"]  ##用于将用户的提示词增强
        self.prompt_keyword = self.prompts["key_words"] ##提取关键词
        self.prompt_abstract = self.prompts["abstract_creator"]
        self.prompt_method = self.prompts["propose_method"]
        self.prompt_paper = self.prompts["paper_writer"]
        self.num_papers = num_papers

    def propose_problem(self,user_input):

        while self.user_satisfied_problem == False:
            completion = client.chat.completions.create(model="gpt-3.5-turbo",
                                            messages=[{"role":"system","content":f"{self.prompt_enhance}"},
                                                      {"role":"user","content":f"{user_input}"},
                                                      ],
                                            )
            enhance_content = completion.choices[0].message.content
            print(enhance_content)
            user_info = input("Do you satisfy the given question?  You can only input yes or no!!\n")
            if user_info == "yes":
                self.user_satisfied_problem = True
            elif user_info == "no":
                self.user_satisfied_problem = False
            else:
                raise("You should input yes or no,program exit!")
            
        self.history = self.history + [enhance_content]
        return enhance_content
    def generate_keywords(self,enhance_content):

        while self.user_satisfied_keywords == False:

            completion = client.chat.completions.create(model="gpt-3.5-turbo",
                                                messages=[{"role":"system","content":f"{self.prompt_keyword}"},
                                                      {"role":"user","content":f"{enhance_content}"},
                                                      ],
                                                )
    
            enhance_keywords = completion.choices[0].message.content
            print(enhance_keywords)
            user_info = input("Do you satisfy the key words?  You can only input yes or no!!\n")
            if user_info == "yes":
                self.user_satisfied_keywords = True
            elif user_info == "no":
                self.user_satisfied_keywords = False
            else:
                raise("You should input yes or no,program exit!")
        self.history = self.history + [enhance_keywords]
        return enhance_keywords
    def find_paper(self,enhance_keywords):
        papers = app.find_basis_paper(enhance_keywords,self.num_papers)
        self.history = self.history + [papers]
        return papers
    def generate_abstract(self,papers):
        completion = client.chat.completions.create(model="gpt-3.5-turbo",
                                                messages=[{"role":"system","content":f"{self.prompt_abstract}"},
                                                      {"role":"user","content":f"papers:{papers}"},
                                                      ],
                                                )
        review = completion.choices[0].message.content
        print(review)   
        self.history = self.history + [review]
        print("\n next is creative method:\n")
        return review   
    def propose_method(self,review):
        while self.user_satisfied_method == False:
            completion = client.chat.completions.create(model="gpt-3.5-turbo",
                                                messages=[{"role":"system","content":f"{self.prompt_method}"},
                                                      {"role":"user","content":f"papers:{review}"},
                                                      ],
                                                    )
            method = completion.choices[0].message.content
            print(method)
            user_info = input("Do you satisfy the method?  You can only input yes or no!!\n")
            if user_info == "yes":
                self.user_satisfied_method = True
                self.history = self.history + [method]
                user_method = input("choose one of the method you satisfied!!\n")
                self.history = self.history + [user_method]
            elif user_info == "no":
                self.user_satisfied_method = False
            else:
                raise("You should input yes or no,program exit!")
        return user_method    
    def generate_papers(self,user_method,papers):
        completion = client.chat.completions.create(model="gpt-3.5-turbo",
                                                messages=[{"role":"system","content":f"{self.prompt_paper}"},
                                                      {"role":"user","content":f"choosen method:{user_method} \n reference paper:{papers}"},
                                                      ],
                                                    )
        paper = completion.choices[0].message.content
        print(paper)
        self.history = self.history + [paper]
        return paper
    
    def return_history(self):
        return self.history
    

his = []
user_input = input("input science topic you are interested")
def paper(user_input):
    word2paper = Word2paper(user_input)
    problem = word2paper.propose_problem(user_input)
    keywords = word2paper.generate_keywords(problem)
    papers = word2paper.find_paper(keywords)
    review = word2paper.generate_abstract(papers)
    method = word2paper.propose_method(review)
    final = word2paper.generate_papers(method,papers)
    his = word2paper.return_history()
    return final,his

final_paper,his = paper(user_input)

##将最终的论文转换为docx文件
def ge_doc():
    doc = Document()
    doc.add_heading(text=his[5])
    text = final_paper
    doc.add_paragraph(text=text)
    doc.save("./output/paper.docx")
    print("文档已保存成功！")

ge_doc()



